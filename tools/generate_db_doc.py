"""
generate_db_doc.py

Create a Word document (`database_schema.docx`) with tables for each database entity.

Usage:
    python tools/generate_db_doc.py [output_path]

Requires:
    pip install python-docx
"""
from docx import Document
from docx.shared import Pt
import sys
from pathlib import Path

SCHEMA = [
    {
        "table": "voters",
        "description": "Voters Table Structure",
        "columns": [
            ("id", "UUID, PRIMARY KEY", "Unique identifier for voter"),
            ("voter_id", "VARCHAR UNIQUE", "Human-readable voter ID (e.g., BVID1234567)"),
            ("aadhaar_id", "VARCHAR ENCRYPTED", "Encrypted Aadhaar number"),
            ("name", "VARCHAR", "Full name of voter"),
            ("email", "VARCHAR UNIQUE", "Email address for OTP"),
            ("phone", "VARCHAR", "Phone number for SMS"),
            ("address", "TEXT", "Residential address"),
            ("state", "VARCHAR", "Indian state"),
            ("constituency", "VARCHAR", "Electoral constituency"),
            ("password_hash", "VARCHAR", "Bcrypt hashed password"),
            ("is_verified", "BOOLEAN", "OTP verification status"),
            ("is_active", "BOOLEAN", "Account active status"),
            ("created_at", "TIMESTAMP", "Registration timestamp"),
            ("updated_at", "TIMESTAMP", "Last update timestamp"),
        ],
    },
    {
        "table": "elections",
        "description": "Elections Table Structure",
        "columns": [
            ("id", "UUID, PRIMARY KEY", "Unique election identifier"),
            ("name", "VARCHAR", "Election name"),
            ("description", "TEXT", "Election details"),
            ("type", "ENUM: General, State, Local", "Election type"),
            ("state", "VARCHAR", "Applicable state (NULL for national)"),
            ("start_date", "TIMESTAMP", "Voting start time"),
            ("end_date", "TIMESTAMP", "Voting end time"),
            ("status", "ENUM: Created, Active, Completed, Cancelled", "Current election status"),
            ("total_votes", "INTEGER", "Total votes cast"),
            ("blockchain_tx_hash", "VARCHAR", "Smart contract transaction hash"),
            ("created_by", "UUID, FK → Admins", "Creating admin reference"),
            ("created_at", "TIMESTAMP", "Creation timestamp"),
            ("updated_at", "TIMESTAMP", "Last modification timestamp"),
        ],
    },
    {
        "table": "votes",
        "description": "Votes Table Structure",
        "columns": [
            ("id", "UUID, PRIMARY KEY", "Unique vote identifier"),
            ("voter_id", "UUID, FK → Voters", "Reference to voter"),
            ("election_id", "UUID, FK → Elections", "Reference to election"),
            ("candidate_id", "UUID, FK → Candidates", "Reference to candidate"),
            ("vote_hash", "VARCHAR UNIQUE", "Cryptographic hash of vote"),
            ("blockchain_tx_hash", "VARCHAR", "Transaction hash on blockchain"),
            ("timestamp", "TIMESTAMP", "Vote casting time"),
            ("is_verified", "BOOLEAN", "Blockchain verification status"),
            ("created_at", "TIMESTAMP", "Record creation time"),
        ],
    },
    {
        "table": "candidates",
        "description": "Candidates Table Structure",
        "columns": [
            ("id", "UUID, PRIMARY KEY", "Unique candidate identifier"),
            ("name", "VARCHAR", "Candidate full name"),
            ("symbol", "VARCHAR", "Party/candidate symbol"),
            ("party", "VARCHAR", "Political party name"),
            ("constituency", "VARCHAR", "Electoral constituency"),
            ("election_id", "UUID, FK → Elections", "Reference to election"),
            ("vote_count", "INTEGER DEFAULT 0", "Total votes received"),
            ("is_verified", "BOOLEAN", "Admin verification status"),
            ("created_at", "TIMESTAMP", "Registration timestamp"),
        ],
    },
    {
        "table": "admins",
        "description": "Admins Table Structure",
        "columns": [
            ("id", "UUID, PRIMARY KEY", "Unique admin identifier"),
            ("email", "VARCHAR UNIQUE", "Admin email address"),
            ("name", "VARCHAR", "Admin full name"),
            ("role", "ENUM: Admin, SuperAdmin", "Admin privilege level"),
            ("state", "VARCHAR", "Assigned state (NULL for super admin)"),
            ("password_hash", "VARCHAR", "Bcrypt hashed password"),
            ("is_active", "BOOLEAN", "Account activation status"),
            ("last_login", "TIMESTAMP", "Last login timestamp"),
            ("created_at", "TIMESTAMP", "Account creation timestamp"),
        ],
    },
    {
        "table": "audit_logs",
        "description": "Audit Logs Table Structure",
        "columns": [
            ("id", "UUID, PRIMARY KEY", "Unique log entry ID"),
            ("actor_id", "UUID, FK → Admins", "Admin who performed action"),
            ("action", "VARCHAR", "Action description (CREATE, UPDATE, DELETE)"),
            ("entity_type", "VARCHAR", "Type of entity modified"),
            ("entity_id", "UUID", "ID of modified entity"),
            ("old_value", "JSON", "Previous value"),
            ("new_value", "JSON", "New value"),
            ("ip_address", "VARCHAR", "Source IP address"),
            ("timestamp", "TIMESTAMP", "Event timestamp"),
        ],
    },
]


def add_table(document: Document, title: str, description: str, columns: list):
    document.add_heading(title, level=2)
    if description:
        document.add_paragraph(description)
    table = document.add_table(rows=1, cols=3)
    table.style = 'Light List Accent 1' if 'Light List Accent 1' in [s.name for s in document.styles] else table.style
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Column Name'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'

    for col_name, col_type, col_desc in columns:
        row_cells = table.add_row().cells
        row_cells[0].text = str(col_name)
        row_cells[1].text = str(col_type)
        row_cells[2].text = str(col_desc)

    document.add_paragraph()


def generate(output_path: Path):
    doc = Document()
    doc.add_heading('Database Schema', level=1)
    doc.add_paragraph('Generated by generate_db_doc.py')

    for entity in SCHEMA:
        add_table(doc, f"{entity['table']}", entity.get('description', ''), entity['columns'])

    doc.save(str(output_path))
    print(f"Saved Word document to: {output_path}")


if __name__ == '__main__':
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('database_schema.docx')
    generate(out)
